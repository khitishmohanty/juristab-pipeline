# src/file_processors.py

from abc import ABC, abstractmethod
from bs4 import BeautifulSoup
import chardet

class FileProcessor(ABC):
    """Abstract base class for file processors"""
    
    @abstractmethod
    def process(self, file_content: bytes, filename: str) -> str:
        """Process the file and return extracted text content"""
        pass

class TxtFileProcessor(FileProcessor):
    """Processor for .txt files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process text files with automatic encoding detection
        """
        try:
            # Try to detect encoding
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            
            # Decode with detected encoding, fallback to utf-8
            try:
                text = file_content.decode(encoding)
            except:
                text = file_content.decode('utf-8', errors='ignore')
            
            # Clean up the text
            text = text.strip()
            
            return text
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")

class HtmlFileProcessor(FileProcessor):
    """Processor for .html files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process HTML files and extract text content
        """
        try:
            # Detect encoding
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            
            # Decode HTML
            try:
                html_text = file_content.decode(encoding)
            except:
                html_text = file_content.decode('utf-8', errors='ignore')
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            raise Exception(f"Error processing HTML file: {str(e)}")

class PdfFileProcessor(FileProcessor):
    """Processor for .pdf files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process PDF files and extract text content
        """
        try:
            import PyPDF2
            import io
            
            # Create a PDF reader object
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except ImportError:
            raise Exception("PyPDF2 library is not installed. Install it with: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")

class DocxFileProcessor(FileProcessor):
    """Processor for .docx files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process DOCX files and extract text content
        """
        try:
            import docx
            import io
            
            # Create a Document object
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            # Extract text from all paragraphs
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        except ImportError:
            raise Exception("python-docx library is not installed. Install it with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {str(e)}")

class CsvFileProcessor(FileProcessor):
    """Processor for .csv files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process CSV files and convert to readable text format
        """
        try:
            import csv
            import io
            
            # Detect encoding
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            
            # Decode CSV
            try:
                csv_text = file_content.decode(encoding)
            except:
                csv_text = file_content.decode('utf-8', errors='ignore')
            
            # Parse CSV
            csv_file = io.StringIO(csv_text)
            csv_reader = csv.reader(csv_file)
            
            # Convert to readable format
            rows = list(csv_reader)
            if not rows:
                return "Empty CSV file"
            
            # Format as text
            text_lines = []
            
            # Add headers if present
            if rows:
                headers = rows[0]
                text_lines.append("CSV Data with columns: " + ", ".join(headers))
                text_lines.append("-" * 50)
                
                # Add data rows
                for i, row in enumerate(rows[1:], 1):
                    row_text = f"Row {i}: "
                    row_items = []
                    for j, (header, value) in enumerate(zip(headers, row)):
                        row_items.append(f"{header}={value}")
                    row_text += ", ".join(row_items)
                    text_lines.append(row_text)
            
            return '\n'.join(text_lines)
        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")

class JsonFileProcessor(FileProcessor):
    """Processor for .json files"""
    
    def process(self, file_content: bytes, filename: str) -> str:
        """
        Process JSON files and convert to readable text format
        """
        try:
            import json
            
            # Detect encoding
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            
            # Decode JSON
            try:
                json_text = file_content.decode(encoding)
            except:
                json_text = file_content.decode('utf-8', errors='ignore')
            
            # Parse JSON
            data = json.loads(json_text)
            
            # Convert to readable format with indentation
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)
            
            return f"JSON Data:\n{formatted_json}"
        except Exception as e:
            raise Exception(f"Error processing JSON file: {str(e)}")

class FileProcessorFactory:
    """Factory class to get the appropriate file processor"""
    
    _processors = {
        'txt': TxtFileProcessor(),
        'html': HtmlFileProcessor(),
        'pdf': PdfFileProcessor(),
        'docx': DocxFileProcessor(),
        'csv': CsvFileProcessor(),
        'json': JsonFileProcessor(),
    }
    
    @classmethod
    def get_processor(cls, file_extension: str) -> FileProcessor:
        """
        Get the appropriate processor for the given file extension
        
        Args:
            file_extension: File extension (without dot)
            
        Returns:
            FileProcessor instance or None if not supported
        """
        return cls._processors.get(file_extension.lower())
    
    @classmethod
    def get_supported_extensions(cls) -> list:
        """Get list of supported file extensions"""
        return list(cls._processors.keys())
    
    @classmethod
    def register_processor(cls, extension: str, processor: FileProcessor):
        """
        Register a new file processor
        
        Args:
            extension: File extension to register
            processor: FileProcessor instance
        """
        cls._processors[extension.lower()] = processor